"""Extraction texte des pieces jointes du chat PALIM (plan P2bis, option A).

Word (.docx), PDF a couche texte (.pdf), Excel (.xlsx/.xlsm), CSV. Un PDF scanne
(couche texte vide) est REFUSE avec un message honnete : l'OCR (Textract) est une
option v2, pas un echec silencieux.

Le texte extrait est plafonne (_MAX_CHARS) : une piece jointe complete un prompt,
elle ne remplace pas l'ingestion RAG.
"""
from __future__ import annotations

import io

_MAX_CHARS = 40_000          # plafond du texte injecte dans le prompt
_MAX_XLSX_CELLS = 4_000      # cellules max converties par classeur
_MAX_FILE_BYTES = 15 * 1024 * 1024

SUPPORTED_TYPES = ["pdf", "docx", "xlsx", "xlsm", "csv"]


class AttachmentError(ValueError):
    """Piece jointe illisible ou non supportee — message destine a l'utilisateur."""


def _truncate(text: str) -> tuple[str, bool]:
    text = text.strip()
    if len(text) <= _MAX_CHARS:
        return text, False
    return text[:_MAX_CHARS] + "\n[... document tronque ...]", True


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n".join(pages).strip()
    if not text:
        raise AttachmentError(
            "Ce PDF semble scanne (aucune couche texte) : je ne peux pas encore le lire. "
            "Transmettez-le au pilote SmarterPlan pour ingestion (OCR)."
        )
    return text


def _extract_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts, cells_used = [], 0
    for ws in wb.worksheets:
        parts.append(f"### Feuille : {ws.title}")
        for row in ws.iter_rows(values_only=True):
            if cells_used >= _MAX_XLSX_CELLS:
                parts.append("[... classeur tronque ...]")
                break
            values = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in values):
                parts.append(" | ".join(values))
                cells_used += len(values)
        if cells_used >= _MAX_XLSX_CELLS:
            break
    wb.close()
    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise AttachmentError("Encodage du CSV non reconnu.")


def extract_attachment(filename: str, data: bytes) -> dict:
    """-> {name, kind, text, truncated}. Leve AttachmentError si illisible/scanne."""
    if len(data) > _MAX_FILE_BYTES:
        raise AttachmentError(f"Fichier trop volumineux (max {_MAX_FILE_BYTES // (1024*1024)} Mo).")
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "docx":
            text = _extract_docx(data)
        elif ext == "pdf":
            text = _extract_pdf(data)
        elif ext in ("xlsx", "xlsm"):
            text = _extract_xlsx(data)
        elif ext == "csv":
            text = _extract_csv(data)
        elif ext == "doc":
            raise AttachmentError("Format .doc (ancien Word) non supporte : enregistrez en .docx.")
        else:
            raise AttachmentError(f"Format .{ext} non supporte ({', '.join(SUPPORTED_TYPES)}).")
    except AttachmentError:
        raise
    except Exception as e:
        raise AttachmentError(f"Lecture impossible de {filename} : {e}") from e
    if not text.strip():
        raise AttachmentError(f"Aucun texte exploitable dans {filename}.")
    text, truncated = _truncate(text)
    return {"name": filename, "kind": ext, "text": text, "truncated": truncated}


def format_for_prompt(att: dict) -> str:
    """Bloc a placer dans le message utilisateur (cf. Bloc 14 des instructions)."""
    note = " (tronque)" if att.get("truncated") else ""
    return (f"[Document joint par l'utilisateur : {att['name']}{note} — piece jointe de "
            f"conversation, PAS un document de la base documentaire]\n"
            f"<<<DEBUT PIECE JOINTE>>>\n{att['text']}\n<<<FIN PIECE JOINTE>>>")
